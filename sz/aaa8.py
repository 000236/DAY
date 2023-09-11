import docx
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
from docx2pdf import convert
import time
import streamlit as st

st.error('请等待右上角的小人跑步结束后输入内容')
text_input = st.text_input("请输入手机号码,输入后请按enter")
t2 = st.text_input("请输入联系人,输入后请按enter")

# 打开Word文件并获取段落列表
doc = docx.Document("sz/2023年第六期大数据师资研修班.docx")
paragraphs = doc.paragraphs

# 将“***”替换为用户输入值a
a = text_input   # 将此处的字符串替换为用户输入的值
b = t2
for p in paragraphs:
    if "*******" in p.text:
        text1 = p.text.replace("*******", a)
        p.clear()
        p.add_run(text1)
for p in paragraphs:
    if "***" in p.text:
        text2 = p.text.replace("***", b)
        p.clear()
        p.add_run(text2)
# 将修改后的Word文件保存到本地
try:
    doc.save("/home/2023年第六期大数据师资研修班_new.docx")
except:
    pass
time.sleep(1)
# 将Word文档转换为PDF
# 将Word文档转换为PDF
def convert_to_pdf(input_path, output_path):
    try:
        convert(input_path, output_path)
        st.write("成功将Word文档转换为PDF！")
    except Exception as e:
        st.write("转换失败:", str(e))

# 指定Word文档路径和输出PDF文件路径
word_file = "/home/2023年第六期大数据师资研修班_new.docx"
pdf_file = "/home/2023年第六期大数据师资研修班_new1.pdf"
convert_to_pdf(word_file, pdf_file)
time.sleep(5)
# 下载重命名后的PDF文件
# if t2 is not None:

with open(pdf_file, "rb") as file:
    file_content = file.read()
    st.download_button("点击下载pdf", data=file_content, file_name=pdf_file)

# # 调用函数进行转换
# st.download_button(
#                 label='下载排行榜',
#                 data=csv,
#                 file_name='排行榜数据1.csv',
#                 key='14'
#             )